import io, time, tempfile
from pathlib import Path

import streamlit as st
from pdf2image import convert_from_bytes
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches
from PIL import Image
import pytesseract

st.set_page_config(page_title="PDF → Word Converter", page_icon="📄", layout="centered")

# ---------------- Utilities ----------------
def compress_image(image: Image.Image, quality: int = 60) -> Image.Image:
    """Nén ảnh JPEG để giảm dung lượng."""
    buf = io.BytesIO()
    image.save(buf, format="JPEG", optimize=True, quality=int(quality))
    buf.seek(0)
    return Image.open(buf)

def pdf_to_word_visual(pdf_bytes: bytes, dpi: int = 150, quality: int = 60, max_pages: int | None = None) -> bytes:
    images = convert_from_bytes(pdf_bytes, dpi=int(dpi))
    if max_pages:
        images = images[:max_pages]

    doc = Document()
    for i, img in enumerate(images):
        img_c = compress_image(img, quality)
        b = io.BytesIO()
        img_c.save(b, format="JPEG")
        b.seek(0)
        if i > 0:
            doc.add_page_break()
        doc.add_picture(b, width=Inches(6.8))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

def pdf_to_word_text(pdf_bytes: bytes, max_pages: int | None = None) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = Path(tmpdir) / "input.pdf"
        out_path = Path(tmpdir) / "output.docx"
        pdf_path.write_bytes(pdf_bytes)

        cv = Converter(str(pdf_path))
        if max_pages:
            cv.convert(str(out_path), start=0, end=max_pages - 1)
        else:
            cv.convert(str(out_path))
        cv.close()

        return out_path.read_bytes()

def pdf_to_word_hybrid(
    pdf_bytes: bytes,
    dpi: int = 150,
    quality: int = 60,
    lang: str = "eng",
    max_pages: int | None = None,
) -> bytes:
    images = convert_from_bytes(pdf_bytes, dpi=int(dpi))
    if max_pages:
        images = images[:max_pages]

    doc = Document()
    for i, img in enumerate(images):
        # chèn ảnh trang
        img_c = compress_image(img, quality)
        b = io.BytesIO()
        img_c.save(b, format="JPEG")
        b.seek(0)
        if i > 0:
            doc.add_page_break()
        doc.add_picture(b, width=Inches(6.8))

        # OCR text
        try:
            text = pytesseract.image_to_string(img, lang=lang).strip()
        except Exception as e:
            text = f"[OCR error: {e}]"
        if text:
            doc.add_paragraph("\n" + text)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

# ---------------- UI ----------------
st.title("📄 PDF → Word")
st.caption("Visual (chèn ảnh), Text (pdf2docx), hoặc Hybrid (ảnh + OCR). Chạy thuần Streamlit.")

with st.sidebar:
    st.header("Tuỳ chọn")
    mode = st.selectbox("Chế độ", ["visual", "text", "hybrid"])
    dpi = st.number_input("DPI (visual/hybrid)", min_value=72, max_value=600, value=150, step=1)
    quality = st.slider("JPEG quality (visual/hybrid)", 1, 95, 60)
    max_pages = st.number_input("Số trang tối đa (0 = tất cả)", min_value=0, value=0, step=1)
    lang = st.text_input("Ngôn ngữ OCR (hybrid)", value="eng")
    st.markdown("Ví dụ: `eng`, `vie`. (Đã chuẩn bị gói tiếng Việt trên Cloud)")

uploaded = st.file_uploader("Chọn file PDF", type=["pdf"])
convert = st.button("Convert")

if convert:
    if not uploaded:
        st.error("Vui lòng chọn file PDF.")
        st.stop()

    pdf_bytes = uploaded.read()
    maxp = int(max_pages) or None
    start = time.time()

    try:
        with st.spinner("Đang chuyển đổi..."):
            if mode == "visual":
                data = pdf_to_word_visual(pdf_bytes, dpi=dpi, quality=quality, max_pages=maxp)
                out_name = uploaded.name.replace(".pdf", "_visual.docx")
            elif mode == "text":
                data = pdf_to_word_text(pdf_bytes, max_pages=maxp)
                out_name = uploaded.name.replace(".pdf", "_text.docx")
            else:
                data = pdf_to_word_hybrid(pdf_bytes, dpi=dpi, quality=quality, lang=lang, max_pages=maxp)
                out_name = uploaded.name.replace(".pdf", "_hybrid.docx")

        st.success(f"Xong trong {time.time() - start:.1f}s")
        st.download_button(
            "⬇️ Tải file Word",
            data=data,
            file_name=out_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.exception(e)
        st.info(
            "Nếu lỗi liên quan Poppler/Tesseract, kiểm tra `packages.txt` (poppler-utils, tesseract-ocr, tesseract-ocr-vie) và redeploy."
        )
